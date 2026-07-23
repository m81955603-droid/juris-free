"""
JURIS-FREE Bolivia — Exportacion de documentos escaneados (multi-pagina)
Combina varias fotos escaneadas en un PDF real, o el texto extraido
de varias paginas en un documento Word editable.

No requiere autenticacion: no toca la base de datos, solo transforma
los datos que el navegador ya tiene (imagenes/texto de la sesion actual).
"""
import base64
import io
import logging
from datetime import datetime
from typing import List, Optional

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches
from fastapi import APIRouter, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.utils import ImageReader
from reportlab.lib import colors

logger = logging.getLogger(__name__)
router = APIRouter()


class Pagina(BaseModel):
    image_base64: str            # imagen JPEG/PNG en base64 (sin el prefijo data:)
    mime_type: str = "image/jpeg"
    texto: Optional[str] = ""    # texto OCR de esta pagina (para exportar a Word)


class ExportRequest(BaseModel):
    paginas: List[Pagina]
    titulo: Optional[str] = "Documento escaneado"


@router.post("/export-pdf")
async def export_pdf(body: ExportRequest):
    """Combina todas las paginas (imagenes) en un solo PDF, una imagen por hoja."""
    if not body.paginas:
        raise HTTPException(400, "No hay paginas para exportar")

    try:
        pdf = fitz.open()
        for pagina in body.paginas:
            img_bytes = base64.b64decode(pagina.image_base64)
            img_doc = fitz.open(stream=img_bytes, filetype="jpg" if "jpeg" in pagina.mime_type or "jpg" in pagina.mime_type else "png")
            rect = img_doc[0].rect
            pdf_page = pdf.new_page(width=rect.width, height=rect.height)
            pdf_page.insert_image(rect, stream=img_bytes)
            img_doc.close()

        buffer = io.BytesIO()
        pdf.save(buffer)
        pdf.close()
        buffer.seek(0)

        return Response(
            content=buffer.read(),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{body.titulo}.pdf"'}
        )
    except Exception as e:
        logger.error(f"Error exportando PDF: {e}")
        raise HTTPException(500, f"No se pudo generar el PDF: {e}")


@router.post("/export-word")
async def export_word(body: ExportRequest):
    """Combina el texto OCR de todas las paginas en un documento Word editable."""
    if not body.paginas:
        raise HTTPException(400, "No hay paginas para exportar")

    try:
        doc = Document()
        doc.add_heading(body.titulo, level=1)

        for i, pagina in enumerate(body.paginas, start=1):
            if len(body.paginas) > 1:
                doc.add_heading(f"Pagina {i}", level=2)
            texto = (pagina.texto or "").strip() or "(sin texto detectado en esta pagina)"
            for parrafo in texto.split("\n"):
                if parrafo.strip():
                    p = doc.add_paragraph(parrafo)
                    p.style.font.size = Pt(11)
            if i < len(body.paginas):
                doc.add_page_break()

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return Response(
            content=buffer.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{body.titulo}.docx"'}
        )
    except Exception as e:
        logger.error(f"Error exportando Word: {e}")
        raise HTTPException(500, f"No se pudo generar el documento Word: {e}")


# ═══════════════════════════════════════════════════════════
# CARNET: anverso + reverso en tamano real, listo para imprimir
# ═══════════════════════════════════════════════════════════

class CarnetPdfRequest(BaseModel):
    front_base64: str             # imagen del anverso, sin el prefijo data:
    back_base64: Optional[str] = None   # imagen del reverso (opcional)
    mime_type: str = "image/jpeg"

# Tamano estandar de carnet/tarjeta ID (ISO/IEC 7810 ID-1), igual al de
# una cedula de identidad o tarjeta de credito: 85.6 x 53.98 mm
CARD_W = 85.6 * mm
CARD_H = 53.98 * mm


CARD_RADIUS = 3 * mm  # esquinas redondeadas, como una tarjeta ID real


def _dibujar_marcas_recorte(c: canvas.Canvas, x: float, y: float, w: float, h: float):
    """Dibuja pequenas lineas guia en las 4 esquinas para recortar con precision."""
    largo = 5 * mm
    gap = 2.5 * mm
    c.setLineWidth(0.5)
    c.setStrokeColor(colors.HexColor("#94a3b8"))

    esquinas = [
        (x, y, -1, -1), (x + w, y, 1, -1),
        (x, y + h, -1, 1), (x + w, y + h, 1, 1),
    ]
    for cx, cy, dx, dy in esquinas:
        c.line(cx + dx*gap, cy, cx + dx*(gap+largo), cy)
        c.line(cx, cy + dy*gap, cx, cy + dy*(gap+largo))


def _dibujar_tarjeta(c: canvas.Canvas, img_bytes: bytes, x: float, y: float, w: float, h: float, etiqueta: str, page_w: float):
    """
    Dibuja una imagen recortada a un rectangulo con esquinas redondeadas
    (como una tarjeta ID real), con borde sutil, sombra suave y etiqueta.
    """
    # Sombra suave (rectangulo gris desplazado, debajo de la tarjeta)
    c.saveState()
    c.setFillColor(colors.HexColor("#e2e8f0"))
    c.roundRect(x + 0.6*mm, y - 0.6*mm, w, h, CARD_RADIUS, fill=1, stroke=0)
    c.restoreState()

    # Imagen recortada a esquinas redondeadas
    c.saveState()
    path = c.beginPath()
    path.roundRect(x, y, w, h, CARD_RADIUS)
    c.clipPath(path, stroke=0, fill=0)
    img = ImageReader(io.BytesIO(img_bytes))
    c.drawImage(img, x, y, width=w, height=h, preserveAspectRatio=True, anchor='c', mask='auto')
    c.restoreState()

    # Borde sutil sobre la tarjeta
    c.saveState()
    c.setStrokeColor(colors.HexColor("#cbd5e1"))
    c.setLineWidth(0.6)
    c.roundRect(x, y, w, h, CARD_RADIUS, fill=0, stroke=1)
    c.restoreState()

    # Etiqueta (ANVERSO / REVERSO)
    c.setFont("Helvetica-Bold", 7.5)
    c.setFillColor(colors.HexColor("#64748b"))
    c.drawCentredString(page_w/2, y + h + 5*mm, etiqueta)

    _dibujar_marcas_recorte(c, x, y, w, h)


@router.post("/export-carnet-pdf")
async def export_carnet_pdf(body: CarnetPdfRequest):
    """
    Genera una hoja A4 con el anverso y reverso del carnet en su
    TAMANO REAL (85.6 x 53.98 mm, igual a una cedula fisica), con
    esquinas redondeadas, sombra suave y lineas guia para recortar.
    Lista para imprimir y laminar.
    """
    try:
        front_bytes = base64.b64decode(body.front_base64)
        back_bytes = base64.b64decode(body.back_base64) if body.back_base64 else None

        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        page_w, page_h = A4

        x = (page_w - CARD_W) / 2
        y_front = page_h - 65*mm - CARD_H
        y_back = 55*mm

        # ── Encabezado con marca ──
        c.setFillColor(colors.HexColor("#1e3a5f"))
        c.roundRect(0, page_h - 18*mm, page_w, 18*mm, 0, fill=1, stroke=0)
        c.setFont("Helvetica-Bold", 13)
        c.setFillColor(colors.white)
        c.drawCentredString(page_w/2, page_h - 11*mm, "MAJA JURÍDICO")
        c.setFont("Helvetica", 8)
        c.setFillColor(colors.HexColor("#cbd5e1"))
        c.drawCentredString(page_w/2, page_h - 15.5*mm, "Documento de identidad — listo para imprimir")

        # ── Tarjetas ──
        _dibujar_tarjeta(c, front_bytes, x, y_front, CARD_W, CARD_H, "ANVERSO", page_w)
        if back_bytes:
            _dibujar_tarjeta(c, back_bytes, x, y_back, CARD_W, CARD_H, "REVERSO", page_w)

        # ── Pie de pagina ──
        fecha = datetime.now().strftime("%d/%m/%Y %H:%M")
        c.setFont("Helvetica", 6.5)
        c.setFillColor(colors.HexColor("#94a3b8"))
        c.drawCentredString(
            page_w/2, 14*mm,
            "Tamaño real de tarjeta ID (85.6 × 53.98 mm) — Recortar por las líneas guía de las esquinas"
        )
        c.drawCentredString(page_w/2, 10*mm, f"Generado el {fecha}")

        c.showPage()
        c.save()
        buffer.seek(0)

        return Response(
            content=buffer.read(),
            media_type="application/pdf",
            headers={"Content-Disposition": 'attachment; filename="carnet_para_imprimir.pdf"'}
        )
    except Exception as e:
        logger.error(f"Error generando PDF de carnet: {e}")
        raise HTTPException(500, f"No se pudo generar el PDF del carnet: {e}")
