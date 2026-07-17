"""
JURIS-FREE Bolivia — Conversor Word <-> PDF
Sin dependencias externas pesadas (no requiere LibreOffice ni MS Word):
usa reportlab para reconstruir PDFs y python-docx + PyMuPDF para leer/
reconstruir documentos Word, preservando texto, negritas/cursivas,
titulos, tablas e imagenes.

Limitacion honesta: para documentos de texto (contratos, memoriales,
resoluciones) el resultado es casi identico al original. Para disenos
graficos complejos (columnas multiples, imagenes en posiciones muy
especificas) las imagenes se colocan en una seccion aparte, no en su
posicion exacta original.
"""
import io
import logging
from typing import List

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import Response
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage, PageBreak
)
from reportlab.lib.enums import TA_JUSTIFY

logger = logging.getLogger(__name__)
router = APIRouter()

MAX_SIZE = 25 * 1024 * 1024  # 25 MB


# ═══════════════════════════════════════════════════════════
# WORD -> PDF
# ═══════════════════════════════════════════════════════════

def _estilo_parrafo(paragraph_style_name: str, styles) -> ParagraphStyle:
    nombre = (paragraph_style_name or "").lower()
    if "heading 1" in nombre or "título 1" in nombre or "titulo 1" in nombre:
        return styles["Heading1"]
    if "heading 2" in nombre or "título 2" in nombre or "titulo 2" in nombre:
        return styles["Heading2"]
    if "heading 3" in nombre or "título 3" in nombre or "titulo 3" in nombre:
        return styles["Heading3"]
    return styles["Justify"]


def _runs_a_html(paragraph) -> str:
    """Convierte los runs de un parrafo docx (negrita/cursiva) a markup reportlab."""
    partes = []
    for run in paragraph.runs:
        texto = (run.text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if not texto:
            continue
        if run.bold:
            texto = f"<b>{texto}</b>"
        if run.italic:
            texto = f"<i>{texto}</i>"
        if run.underline:
            texto = f"<u>{texto}</u>"
        partes.append(texto)
    return "".join(partes) or (paragraph.text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _extraer_imagenes_docx(doc: DocxDocument) -> List[bytes]:
    imagenes = []
    try:
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                imagenes.append(rel.target_part.blob)
    except Exception as e:
        logger.warning(f"No se pudieron extraer imagenes del docx: {e}")
    return imagenes


@router.post("/word-to-pdf")
async def word_to_pdf(file: UploadFile = File(...)):
    if not file.filename.lower().endswith((".docx",)):
        raise HTTPException(400, "Solo se aceptan archivos .docx")

    contenido = await file.read()
    if len(contenido) > MAX_SIZE:
        raise HTTPException(400, "El archivo supera el límite de 25 MB")

    try:
        doc = DocxDocument(io.BytesIO(contenido))

        buffer = io.BytesIO()
        pdf_doc = SimpleDocTemplate(
            buffer, pagesize=LETTER,
            leftMargin=1*inch, rightMargin=1*inch, topMargin=1*inch, bottomMargin=1*inch
        )
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name="Justify", parent=styles["Normal"], alignment=TA_JUSTIFY, spaceAfter=8, fontSize=10.5, leading=15))

        elementos = []

        for para in doc.paragraphs:
            texto_html = _runs_a_html(para)
            if not texto_html.strip():
                elementos.append(Spacer(1, 8))
                continue
            estilo = _estilo_parrafo(para.style.name if para.style else "", styles)
            try:
                elementos.append(Paragraph(texto_html, estilo))
            except Exception:
                # Si el markup falla (caracteres raros), usar texto plano
                elementos.append(Paragraph(para.text, styles["Justify"]))

        for tabla in doc.tables:
            datos = [[celda.text for celda in fila.cells] for fila in tabla.rows]
            if not datos:
                continue
            t = Table(datos, hAlign="LEFT")
            t.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f1f5f9")),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ]))
            elementos.append(Spacer(1, 10))
            elementos.append(t)
            elementos.append(Spacer(1, 10))

        imagenes = _extraer_imagenes_docx(doc)
        if imagenes:
            elementos.append(PageBreak())
            elementos.append(Paragraph("Imágenes del documento", styles["Heading2"]))
            for img_bytes in imagenes:
                try:
                    img_buf = io.BytesIO(img_bytes)
                    rl_img = RLImage(img_buf, width=5*inch, height=3.5*inch, kind="proportional")
                    elementos.append(rl_img)
                    elementos.append(Spacer(1, 12))
                except Exception:
                    continue

        if not elementos:
            elementos.append(Paragraph("(Documento vacío)", styles["Normal"]))

        pdf_doc.build(elementos)
        buffer.seek(0)

        nombre_salida = file.filename.rsplit(".", 1)[0] + ".pdf"
        return Response(
            content=buffer.read(),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{nombre_salida}"'}
        )
    except Exception as e:
        logger.error(f"Error convirtiendo Word a PDF: {e}")
        raise HTTPException(500, f"No se pudo convertir el documento: {e}")


# ═══════════════════════════════════════════════════════════
# PDF -> WORD
# ═══════════════════════════════════════════════════════════

def _es_negrita(flags: int) -> bool:
    return bool(flags & 2**4)


def _es_italica(flags: int) -> bool:
    return bool(flags & 2**1)


@router.post("/pdf-to-word")
async def pdf_to_word(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(400, "Solo se aceptan archivos .pdf")

    contenido = await file.read()
    if len(contenido) > MAX_SIZE:
        raise HTTPException(400, "El archivo supera el límite de 25 MB")

    try:
        pdf = fitz.open(stream=contenido, filetype="pdf")
        doc = DocxDocument()

        for num_pagina, pagina in enumerate(pdf):
            if num_pagina > 0:
                doc.add_page_break()

            texto_dict = pagina.get_text("dict")
            hubo_contenido = False

            for bloque in texto_dict.get("blocks", []):
                if bloque.get("type") != 0:  # 0 = texto, 1 = imagen
                    continue
                for linea in bloque.get("lines", []):
                    parrafo_docx = doc.add_paragraph()
                    for span in linea.get("spans", []):
                        texto = span.get("text", "")
                        if not texto.strip():
                            continue
                        hubo_contenido = True
                        run = parrafo_docx.add_run(texto)
                        tam = span.get("size", 11)
                        run.font.size = Pt(max(7, min(28, round(tam))))
                        flags = span.get("flags", 0)
                        run.font.bold = _es_negrita(flags)
                        run.font.italic = _es_italica(flags)

            # Extraer imagenes de la pagina y agregarlas al final de esa seccion
            for img_index, img in enumerate(pagina.get_images(full=True)):
                try:
                    xref = img[0]
                    base_img = pdf.extract_image(xref)
                    img_bytes = base_img["image"]
                    img_stream = io.BytesIO(img_bytes)
                    doc.add_picture(img_stream, width=Inches(5.5))
                    hubo_contenido = True
                except Exception:
                    continue

            if not hubo_contenido:
                doc.add_paragraph(f"(Página {num_pagina + 1} sin texto extraíble — puede ser una imagen escaneada)")

        pdf.close()

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        nombre_salida = file.filename.rsplit(".", 1)[0] + ".docx"
        return Response(
            content=buffer.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{nombre_salida}"'}
        )
    except Exception as e:
        logger.error(f"Error convirtiendo PDF a Word: {e}")
        raise HTTPException(500, f"No se pudo convertir el documento: {e}")
